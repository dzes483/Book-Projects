#!/usr/bin/python3
# custom_invitations.py - Creates custom invitations from a .txt file of guests
# and a pre-styled .docx document.

import docx

# Open already-styled invitation
doc = input('Enter the name of the styled template file: ')
template = docx.Document(doc)

# Open guest list
guests = input('Enter the name of the guest list .txt file: ')
guest_list = open(guests, 'r')

# Define styles
styles = template.styles

# Loop through the guest list, adding the default text along with the guest's
# name
print('Printing invitations...')
for i in guest_list:
    intro = template.add_paragraph('It would be a pleasure to have the company\
                                     of\n', style='Heading1')
    guest_name = template.add_paragraph(i, style='Heading2')
    address = template.add_paragraph('at 11010 Memory Lane on the evening of',
                                     style='Heading1')
    date = template.add_paragraph('April 1st', style='Heading3')
    time = template.add_paragraph("at 7 o'clock", style='Heading1')
    page_break = template.add_page_break()

print('Saving as completed_invitations.docx...')
template.save('completed_invitations.docx')
print('Done.')
